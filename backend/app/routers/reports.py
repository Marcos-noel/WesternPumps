from __future__ import annotations

import csv
import io
from datetime import datetime, timedelta

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.db import get_db
from app.deps import require_roles
from app.models import AuditLog, ItemInstance, Part, PartLocationStock, StockTransaction, UsageRecord

try:
    from docx import Document
except Exception:  # pragma: no cover - handled at runtime
    Document = None


router = APIRouter(prefix="/api/reports", tags=["reports"])


def _format_response(data: bytes, filename: str, content_type: str) -> Response:
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(content=data, media_type=content_type, headers=headers)


def _build_excel(headers: list[str], rows: list[list[str]], title: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = title
    ws.append(headers)
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_csv(headers: list[str], rows: list[list[str]]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(headers)
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


def _build_pdf(title: str, headers: list[str], rows: list[list[str]]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, 770, title)
    c.setFont("Helvetica", 9)
    y = 740
    c.drawString(40, y, " | ".join(headers))
    y -= 16
    for row in rows:
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 9)
            y = 770
        c.drawString(40, y, " | ".join(row)[:180])
        y -= 13
    c.save()
    return buf.getvalue()


def _build_docx(title: str, headers: list[str], rows: list[list[str]]) -> bytes:
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx not installed")
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _format_export(
    *,
    fmt: str,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    basename: str,
) -> Response:
    if fmt == "excel":
        data = _build_excel(headers, rows, title)
        return _format_response(data, f"{basename}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    if fmt == "csv":
        data = _build_csv(headers, rows)
        return _format_response(data, f"{basename}.csv", "text/csv")
    if fmt == "pdf":
        data = _build_pdf(title, headers, rows)
        return _format_response(data, f"{basename}.pdf", "application/pdf")
    if fmt == "docx":
        data = _build_docx(title, headers, rows)
        return _format_response(data, f"{basename}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    raise HTTPException(status_code=400, detail="Invalid format")


@router.get("/stock-level", dependencies=[Depends(require_roles("manager", "finance"))])
def stock_level_report(
    db: Session = Depends(get_db),
    format: str = Query("excel", pattern="^(excel|pdf|docx|csv)$"),
    q: str | None = Query(None, max_length=200),
    low_only: bool = Query(False),
    part_id: int | None = Query(None),
    category_id: int | None = Query(None),
    location_id: int | None = Query(None),
    limit: int = Query(1000, ge=1, le=10000),
) -> Response:
    stmt = select(Part).order_by(Part.name.asc())
    if q:
        like = f"%{q.strip()}%"
        stmt = stmt.where((Part.sku.like(like)) | (Part.name.like(like)))
    if low_only:
        stmt = stmt.where(Part.quantity_on_hand <= Part.min_quantity)
    if part_id is not None:
        stmt = stmt.where(Part.id == part_id)
    if category_id is not None:
        stmt = stmt.where(Part.category_id == category_id)
    if location_id is not None:
        location_part_ids = select(PartLocationStock.part_id).where(PartLocationStock.location_id == location_id)
        stmt = stmt.where((Part.location_id == location_id) | (Part.id.in_(location_part_ids)))

    parts = db.scalars(stmt.limit(limit)).all()
    headers = ["SKU", "Barcode", "Name", "Tracking", "Qty On Hand", "Min Qty", "Unit Cost", "Category ID"]
    rows = [
        [
            p.sku,
            p.barcode_value or "",
            p.name,
            p.tracking_type,
            str(p.quantity_on_hand),
            str(p.min_quantity),
            f"{p.unit_price or 0:.2f}" if p.unit_price is not None else "",
            str(p.category_id) if p.category_id else "",
        ]
        for p in parts
    ]

    return _format_export(
        fmt=format,
        title="Stock Levels Report",
        headers=headers,
        rows=rows,
        basename="stock-levels",
    )


@router.get("/item-traceability", dependencies=[Depends(require_roles("manager", "finance"))])
def item_traceability_report(
    db: Session = Depends(get_db),
    format: str = Query("excel", pattern="^(excel|pdf|docx|csv)$"),
    serial_number: str | None = Query(None, max_length=100),
    item_instance_id: int | None = Query(None),
    limit: int = Query(1000, ge=1, le=10000),
) -> Response:
    if not serial_number and item_instance_id is None:
        raise HTTPException(status_code=400, detail="serial_number or item_instance_id is required")

    item_stmt = select(ItemInstance).order_by(ItemInstance.id.asc())
    if item_instance_id is not None:
        item_stmt = item_stmt.where(ItemInstance.id == item_instance_id)
    if serial_number:
        item_stmt = item_stmt.where(ItemInstance.serial_number == serial_number.strip())
    instances = db.scalars(item_stmt.limit(1)).all()
    if not instances:
        raise HTTPException(status_code=404, detail="Item instance not found")
    instance = instances[0]

    tx_stmt = (
        select(StockTransaction)
        .where(StockTransaction.item_instance_id == instance.id)
        .order_by(StockTransaction.created_at.asc())
        .limit(limit)
    )
    usage_stmt = (
        select(UsageRecord)
        .where(UsageRecord.item_instance_id == instance.id)
        .order_by(UsageRecord.used_at.asc())
        .limit(limit)
    )
    txs = db.scalars(tx_stmt).all()
    usage_rows = db.scalars(usage_stmt).all()

    headers = [
        "Timestamp",
        "Event Type",
        "Part ID",
        "Serial Number",
        "Barcode Value",
        "Transaction Type",
        "Movement",
        "Quantity Delta",
        "Request ID",
        "Technician ID",
        "Customer ID",
        "Job ID",
        "Reference",
    ]
    rows: list[list[str]] = []
    for t in txs:
        rows.append(
            [
                t.created_at.isoformat(sep=" ", timespec="seconds"),
                "STOCK_TRANSACTION",
                str(t.part_id),
                instance.serial_number,
                instance.barcode_value or "",
                t.transaction_type.value,
                t.movement_type or "",
                str(t.quantity_delta),
                str(t.request_id) if t.request_id else "",
                str(t.technician_id) if t.technician_id else "",
                str(t.customer_id) if t.customer_id else "",
                str(t.job_id) if t.job_id else "",
                f"tx:{t.id}",
            ]
        )
    for u in usage_rows:
        rows.append(
            [
                u.used_at.isoformat(sep=" ", timespec="seconds"),
                "USAGE_RECORD",
                str(instance.part_id),
                instance.serial_number,
                instance.barcode_value or "",
                "",
                "USAGE",
                "",
                str(u.request_id) if u.request_id else "",
                str(u.technician_id),
                str(u.customer_id) if u.customer_id else "",
                str(u.job_id) if u.job_id else "",
                f"usage:{u.id}",
            ]
        )
    rows.sort(key=lambda r: r[0])

    return _format_export(
        fmt=format,
        title="Item Traceability Report",
        headers=headers,
        rows=rows,
        basename=f"item-traceability-{instance.serial_number}",
    )


@router.get("/stock-movement", dependencies=[Depends(require_roles("manager", "finance"))])
def stock_movement_report(
    db: Session = Depends(get_db),
    format: str = Query("excel", pattern="^(excel|pdf|docx|csv)$"),
    start: datetime | None = Query(None),
    end: datetime | None = Query(None),
    transaction_type: str | None = Query(None, max_length=10),
    part_id: int | None = Query(None),
    technician_id: int | None = Query(None),
    limit: int = Query(2000, ge=1, le=20000),
) -> Response:
    stmt = select(StockTransaction).order_by(StockTransaction.created_at.desc())
    if start:
        stmt = stmt.where(StockTransaction.created_at >= start)
    if end:
        stmt = stmt.where(StockTransaction.created_at <= end)
    if transaction_type:
        stmt = stmt.where(StockTransaction.transaction_type == transaction_type)
    if part_id is not None:
        stmt = stmt.where(StockTransaction.part_id == part_id)
    if technician_id is not None:
        stmt = stmt.where(StockTransaction.technician_id == technician_id)

    txs = db.scalars(stmt.limit(limit)).all()
    headers = ["Date", "Part ID", "Type", "Delta", "Movement", "Request", "Technician", "Created By"]
    rows: list[list[str]] = []
    for t in txs:
        rows.append(
            [
                t.created_at.isoformat(sep=" ", timespec="seconds"),
                str(t.part_id),
                t.transaction_type.value,
                str(t.quantity_delta),
                t.movement_type or "",
                str(t.request_id) if t.request_id else "",
                str(t.technician_id) if t.technician_id else "",
                str(t.created_by_user_id) if t.created_by_user_id else "",
            ]
        )

    return _format_export(
        fmt=format,
        title="Stock Movement Report",
        headers=headers,
        rows=rows,
        basename="stock-movements",
    )


@router.get("/audit-trail", dependencies=[Depends(require_roles("manager", "finance"))])
def audit_trail_report(
    db: Session = Depends(get_db),
    format: str = Query("excel", pattern="^(excel|pdf|docx|csv)$"),
    start: datetime | None = Query(None),
    end: datetime | None = Query(None),
    user_id: int | None = Query(None),
    entity_type: str | None = Query(None, max_length=100),
    action: str | None = Query(None, max_length=100),
    limit: int = Query(5000, ge=1, le=50000),
) -> Response:
    stmt = select(AuditLog).order_by(AuditLog.created_at.desc())
    if start:
        stmt = stmt.where(AuditLog.created_at >= start)
    if end:
        stmt = stmt.where(AuditLog.created_at <= end)
    if user_id is not None:
        stmt = stmt.where(AuditLog.user_id == user_id)
    if entity_type:
        stmt = stmt.where(AuditLog.entity_type == entity_type)
    if action:
        stmt = stmt.where(AuditLog.action == action)

    logs = db.scalars(stmt.limit(limit)).all()
    headers = ["Timestamp", "User ID", "Action", "Entity Type", "Entity ID", "Prev Hash", "Entry Hash", "Detail"]
    rows = [
        [
            l.created_at.isoformat(sep=" ", timespec="seconds"),
            str(l.user_id) if l.user_id else "",
            l.action,
            l.entity_type,
            str(l.entity_id) if l.entity_id is not None else "",
            l.prev_hash or "",
            l.entry_hash or "",
            l.detail or "",
        ]
        for l in logs
    ]

    return _format_export(
        fmt=format,
        title="Audit Trail Report",
        headers=headers,
        rows=rows,
        basename="audit-trail",
    )


@router.get("/audit-integrity", dependencies=[Depends(require_roles("manager", "finance"))])
def audit_integrity_report(
    db: Session = Depends(get_db),
    limit: int = Query(50000, ge=1, le=200000),
) -> dict[str, object]:
    logs = db.scalars(select(AuditLog).order_by(AuditLog.id.asc()).limit(limit)).all()
    broken_ids: list[int] = []
    previous_hash: str | None = None

    import hashlib

    for log in logs:
        if (log.prev_hash or None) != (previous_hash or None):
            broken_ids.append(log.id)
        basis = "|".join(
            [
                str(log.user_id if log.user_id is not None else ""),
                log.action,
                log.entity_type,
                str(log.entity_id if log.entity_id is not None else ""),
                log.detail or "",
                log.prev_hash or "",
            ]
        )
        expected = hashlib.sha256(basis.encode("utf-8")).hexdigest()
        if (log.entry_hash or "") != expected:
            broken_ids.append(log.id)
        previous_hash = log.entry_hash

    return {
        "checked": len(logs),
        "valid": len(broken_ids) == 0,
        "broken_ids": sorted(set(broken_ids))[:200],
    }


@router.get("/forecast", dependencies=[Depends(require_roles("manager", "finance"))])
def stock_forecast(
    db: Session = Depends(get_db),
    days: int = Query(30, ge=7, le=365),
    lookback_days: int = Query(30, ge=7, le=365),
) -> dict[str, object]:
    """
    Lightweight forecast using daily net stock movement average.
    Positive => expected inflow, negative => expected outflow.
    """
    today = datetime.utcnow().date()
    start = today - timedelta(days=lookback_days - 1)

    txs = db.scalars(select(StockTransaction).where(StockTransaction.created_at >= datetime.combine(start, datetime.min.time()))).all()
    by_day: dict[str, int] = {}
    for tx in txs:
        key = tx.created_at.date().isoformat()
        by_day[key] = by_day.get(key, 0) + int(tx.quantity_delta or 0)

    observed_series: list[dict[str, object]] = []
    total_net = 0
    cursor = start
    while cursor <= today:
        key = cursor.isoformat()
        value = int(by_day.get(key, 0))
        total_net += value
        observed_series.append({"date": key, "net": value})
        cursor += timedelta(days=1)

    avg_daily_net = total_net / float(lookback_days)
    forecast_series: list[dict[str, object]] = []
    future = today + timedelta(days=1)
    for _ in range(days):
        forecast_series.append({"date": future.isoformat(), "projected_net": round(avg_daily_net, 2)})
        future += timedelta(days=1)

    return {
        "method": "rolling_average_daily_net",
        "lookback_days": lookback_days,
        "forecast_days": days,
        "average_daily_net": round(avg_daily_net, 2),
        "observed": observed_series,
        "forecast": forecast_series,
    }
