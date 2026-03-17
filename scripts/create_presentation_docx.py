from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create document
doc = Document()

# Title
title = doc.add_heading('WesternPumps Inventory Management System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Complete Feature Overview for Client Presentations')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'WesternPumps is a comprehensive, enterprise-grade inventory and business management system designed for medium to large organizations. Built with modern technology (React + FastAPI), it provides a complete solution for managing customers, jobs, inventory, and deliveries with built-in audit trails, real-time monitoring, and scalability for high-volume operations.'
)

# Key Statistics
doc.add_heading('Key Statistics:', level=2)
stats = [
    'Built on FastAPI (used by Netflix, Uber, Microsoft)',
    'Supports 10,000+ requests per day',
    '99.9% availability target (SLA-backed)',
    'Real-time performance monitoring included'
]
for stat in stats:
    doc.add_paragraph(stat, style='List Bullet')

# Core Modules
doc.add_heading('2. Core Modules', level=1)

doc.add_heading('2.1 Customer Management', level=2)
doc.add_paragraph('Full customer directory with contact information, fast search functionality, and link to service jobs.')

doc.add_heading('2.2 Job Tracking', level=2)
doc.add_paragraph('Service job management with statuses (Open, In Progress, Completed, Cancelled), priority levels, and staff assignment.')

doc.add_heading('2.3 Inventory Management', level=2)
doc.add_paragraph('Items/Parts Management, Supplier Management, and Stock Transactions with complete audit trail.')

doc.add_heading('2.4 Stock Requests & Workflow', level=2)
doc.add_paragraph('Staff can request parts with multi-level approval workflow and value-based thresholds.')

doc.add_heading('2.5 Deliveries', level=2)
doc.add_paragraph('Delivery management with status tracking and customer notification integration.')

# Standout Features
doc.add_heading('3. Standout Features (Why Choose WesternPumps)', level=1)

doc.add_heading('3.1 Audit-First Inventory Model', level=2)
doc.add_paragraph('Most inventory systems just show current quantities - you cannot explain why stock changed. WesternPumps uses dual-track system: Current on-hand quantity (fast reads) + Complete transaction ledger (IN/OUT/ADJUST). Benefit: Never wonder "where did my stock go?" again.')

doc.add_heading('3.2 Fast Data Entry', level=2)
doc.add_paragraph('Keyboard-friendly design with minimal clicks, quick search and filters, and fast customer lookup.')

doc.add_heading('3.3 Built-in AI Assistant', level=2)
doc.add_paragraph('Smart queries about inventory, business insights, automation recommendations, and natural language interface.')

doc.add_heading('3.4 Role-Based Access Control', level=2)
doc.add_paragraph('Admin, Store Manager, Technician, Manager, and Approver roles with specific permissions.')

doc.add_heading('3.5 Approval Workflows', level=2)
doc.add_paragraph('Configurable approval thresholds based on value with complete audit trail.')

# User Roles and Capabilities
doc.add_heading('4. User Roles & Capabilities', level=1)

doc.add_paragraph('WesternPumps has 10 user roles with specific capabilities for granular access control:')

roles_table = doc.add_table(rows=11, cols=2)
roles_table.style = 'Table Grid'

role_headers = roles_table.rows[0].cells
role_headers[0].text = 'Role'
role_headers[1].text = 'Capabilities'

roles_data = [
    ('admin', 'Full system access, user management, create admins, all API endpoints'),
    ('manager', 'Approval and reporting, view all data, approve stock requests'),
    ('store_manager', 'Inventory custody, receive/issue stock, manage transactions'),
    ('lead_technician', 'Technician capabilities plus can assign jobs to technicians'),
    ('technician', 'Request parts, record usage, create/update assigned jobs'),
    ('staff', '(Legacy - mapped to technician) Same as technician'),
    ('approver', 'Approve/reject stock requests, update stock quantities'),
    ('finance', 'View reports, generate reports, set reorder thresholds'),
    ('rider', 'Delivery-related tasks'),
    ('driver', 'Delivery-related tasks')
]

for i, (role, caps) in enumerate(roles_data):
    row = roles_table.rows[i+1].cells
    row[0].text = role
    row[1].text = caps

doc.add_heading('AI Assistant Role Permissions:', level=2)
doc.add_paragraph('Create Product: admin, manager, store_manager')
doc.add_paragraph('Update Stock: admin, manager, store_manager, approver')
doc.add_paragraph('Delete Product: admin, manager, store_manager, approver')
doc.add_paragraph('Generate Report: admin, manager, finance')
doc.add_paragraph('Bulk Update: admin, manager, store_manager, approver')
doc.add_paragraph('Set Reorder: manager, finance')

# Technical Excellence
doc.add_heading('5. Technical Excellence', level=1)

doc.add_heading('5.1 Performance and Scalability', level=2)
doc.add_paragraph('FastAPI Backend (one of the fastest Python frameworks), Async/Await for concurrent requests, Horizontal Scaling, Connection Pooling, and 10,000+ requests/day easily handled.')

doc.add_heading('5.2 Built-in Monitoring', level=2)
doc.add_paragraph('Real-time metrics tracking, P95 latency monitoring, 99.9% availability SLA, Prometheus integration, and health check endpoints.')

doc.add_heading('5.3 Security Features', level=2)
doc.add_paragraph('JWT authentication, role-based permissions, security headers (CSP, HSTS), HTTPS enforcement, request logging, and audit trails.')

# Enterprise Features
doc.add_heading('6. Enterprise Features', level=1)

doc.add_paragraph('Integration Capabilities: Finance, ERP, Accounting')
doc.add_paragraph('Observability: Prometheus, Grafana, OpenTelemetry')
doc.add_paragraph('OIDC/SSO Ready for enterprise authentication')
doc.add_paragraph('Multi-Tenant Architecture')
doc.add_paragraph('Event-Driven Design with WebSocket support')

# Why WesternPumps for Kenyan Businesses
doc.add_heading('7. Why WesternPumps for Kenyan Businesses?', level=1)

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'

headers = table.rows[0].cells
headers[0].text = 'Feature'
headers[1].text = 'Benefit for Kenyan Market'

data = [
    ('Self-Hosted', 'No recurring SaaS fees - own your data'),
    ('Audit Trail', 'Solves inventory shrinkage problem'),
    ('Offline Ready', 'SQLite option for areas with poor connectivity'),
    ('Fast Entry', 'High-volume operations made efficient'),
    ('Local Support', 'Built for East African business needs'),
    ('Multi-Currency', 'Ready for KES and USD')
]

for i, (feature, benefit) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = feature
    row[1].text = benefit

# Comparison
doc.add_heading('8. Comparison with Typical Kenyan Inventory Systems', level=1)

comp_table = doc.add_table(rows=8, cols=3)
comp_table.style = 'Table Grid'

comp_headers = comp_table.rows[0].cells
comp_headers[0].text = 'Feature'
comp_headers[1].text = 'WesternPumps'
comp_headers[2].text = 'Typical Kenyan Solutions'

comparison_data = [
    ('Audit Trail', 'Full', 'None'),
    ('AI Assistant', 'Built-in', 'Not available'),
    ('Performance', '10k+ requests', 'Limited'),
    ('Scalability', 'Horizontal', 'Limited'),
    ('Monitoring', 'Built-in', 'None'),
    ('Integration', 'Multiple APIs', 'Limited'),
    ('Self-Hosted', 'Yes', 'Usually SaaS only')
]

for i, (feature, wp, typical) in enumerate(comparison_data):
    row = comp_table.rows[i+1].cells
    row[0].text = feature
    row[1].text = wp
    row[2].text = typical

# Summary
doc.add_heading('9. Summary', level=1)
doc.add_paragraph('WesternPumps is not just another inventory system - it is a complete business management platform with:')
doc.add_paragraph('1. Complete Visibility - Know exactly what you have and where it went')
doc.add_paragraph('2. Audit Confidence - Every transaction is traceable')
doc.add_paragraph('3. Performance - Handles high-volume operations')
doc.add_paragraph('4. Enterprise Ready - Scales with your business')
doc.add_paragraph('5. Cost Effective - Self-hosted, no recurring fees')
doc.add_paragraph('6. Modern Technology - Fast, reliable, maintainable')

doc.add_heading('Ideal for:', level=2)
doc.add_paragraph('Service companies with inventory')
doc.add_paragraph('Distributors and wholesalers')
doc.add_paragraph('Manufacturing businesses')
doc.add_paragraph('Any organization needing audit-ready inventory')

# Save
doc.save('docs/WesternPumps_Presentation_v2.docx')
print("DOCX file created successfully!")
